import openpyxl
import pandas as pd
from pyautogui import size
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import subprocess
import shutil
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
from bs4 import BeautifulSoup
import time
import datetime
import pyautogui
import pyperclip
import csv
import sys
import os
import math
import requests
import re
import random
import chromedriver_autoinstaller
from PyQt5.QtWidgets import QWidget, QApplication, QTreeView, QFileSystemModel, QVBoxLayout, QPushButton, QInputDialog, \
    QLineEdit, QMainWindow, QMessageBox, QFileDialog
from PyQt5.QtCore import QCoreApplication
from selenium.webdriver import ActionChains
from datetime import datetime, date, timedelta
import numpy
import datetime
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
import json
import pprint
from docx import Document
from docx.shared import RGBColor
import docx
from docx import Document
from docx.shared import RGBColor, Pt
from window import Ui_MainWindow
from tkinter import Tk
import firebase_admin
from firebase_admin import credentials
from firebase_admin import db
import copy
import socket
from collections import deque


def get_key(pc_no,first_flag):
    if first_flag==True:
        secret_key={
          "type": "service_account",
          "project_id": "bigkinds-8d2eb",
          "private_key_id": "f03db71c1a6f2e56ee6963d1b3605c2a65d9f358",
          "private_key": "-----BEGIN PRIVATE KEY-----\nMIIEvQIBADANBgkqhkiG9w0BAQEFAASCBKcwggSjAgEAAoIBAQDQxsq11RjOdxq6\nQCDPZ/ZMgJkBSjwF+y/c3cn/IGZBpuUbRCWIH5sETPj1EZKNNQq1tgiJmAVJmBrb\n2MpuE1pecSah1QZTC/isms8Sbjd8PQTaPOEKaxe8caEUlJB/qP9u+XlPl4/F1VtC\n1FFAeGX1VTaNUHvFD+Bx3ZqPTBDsxTbLCpvP2NFm3uUaTvfVPx5CT/GFrx43Q3yh\naTstnTVih3P5wdW8jpoPREivCYr+55i28ePZaXp3yUDtbPpQ/ViWq0jYNQftejxb\nMMFpDia7puwVUlahmyVKtOuhyGkwYHOceucArfA2OOGGsrsLi7fPMWApCHSjqm/q\nqwvv1NlBAgMBAAECggEAFHtkOLmKwPqrWxlefqqO3cuB9z/uqpGNjQgPQKHPlmJG\nuMmTOwsKvL3GcNiD/al1/wS4SrNo1WMAQ5A9r9XjOD+2kArjn3MdtuVLKl44wE4G\nM1jkpQ6Dmh+1s1yt3nBYeu0rqxTN7JU1alWFESOesRZmjy+VrOwVU2Zvut42DNxQ\nuoTj5f4ipAbLV9h3JS2dN0eZLH4bktD+rRDzbmhZ9Nog9XphRwpcNQcI+PWB0xaD\nMeG7BDwOyLPRGPl8/hlN50tqDihgGxqIJdoDBodZ2aYLuhNQpBK1H/FWA8hDmDPz\nchGzjTJQ0jlbwSUnCjYmVi/Mdk35TT6fOhsuAQdLrwKBgQD5EWtQyBdDmLkM3K/8\n5QqFN7vJJDn4IgfxY/nL9W0jshUBoC/EXnlmYf/hK6nLeAzszpVekD4VEZ2RW737\nGvsFcvo6brqWcgiss7YlUz6dEoZgXlcgX0xNdN2me1Uyddt+Ob9MmaYLL1kXD0re\nCf5z5oNGTczt+/NSNBfzWQL1MwKBgQDWlkzgU6oyPagMu7AFpKnNQ81qaaYKbvTP\nbFqcTB/jnQ2URO6hhrXFlriwN+H/Fwk+2sRdtaqb5cIuHTeMfwpjM5DUsfeJQONT\norHR9IyIDFBo/W9T/IkgZRUMDHVtslC7uwo3SI5+Vq3cpHtOaKYmDZWl07rLjID3\ngAy2CPtPuwKBgDxQCYqwSWHnE7iXoE4/MIL98g1NJnR92SqNKSpIrjscnpWcMrT8\no601Qf9G0YEh+w9FH0qc50u2zk8qxiCOcgbpxsprkedoic04X30+YUUr/TpM3hBK\nw3FmQV5VM1Zaz+ILHdXwxxKiTpKU8249QU7TnSpjIKLohg176+jTvzGdAoGAUqmn\nh6rzH1N9DxqmHiZmevopgVqdaduAVv8okJkKl8YCWyGKw5J/J1R+ZPPCDWA0YZAo\noOjAhjml6dm7clXmDHzqx+SetLLYX3mHApgwIqCLwUYPwsy/jkkiKHgMLGvKiebt\ncopyydxeMWbzUSsjBfmpsi20AYFolO8w9lbTaM0CgYEA+LTi6LaFJEG9JMA0xixQ\nQ6M5ZtHZnhhO0o4fI/KKmcua/vNwAHUAt7JvYVgAcj54x5CEMSw7skM30Lz+hnsB\nBupTgUwxqL8aSPLL2Cv/naf5SJoDgBsmUMCsUwwweLdTHxELY+1UE47BCv0oyC5a\nLdvgLjFM1uyQ1XXlAPiXE+s=\n-----END PRIVATE KEY-----\n",
          "client_email": "firebase-adminsdk-4sjyo@bigkinds-8d2eb.iam.gserviceaccount.com",
          "client_id": "116575592744504231900",
          "auth_uri": "https://accounts.google.com/o/oauth2/auth",
          "token_uri": "https://oauth2.googleapis.com/token",
          "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
          "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/firebase-adminsdk-4sjyo%40bigkinds-8d2eb.iam.gserviceaccount.com"
        }

        cred=credentials.Certificate(secret_key)
        firebase_admin.initialize_app(cred,{
            'databaseURL':"https://bigkinds-8d2eb-default-rtdb.asia-southeast1.firebasedatabase.app/"
        })

    ref=db.reference().get()
    result_password=ref['password'][pc_no]['password']
    result_ip = ref['password'][pc_no]['ip']

    return result_password, result_ip
def make_word(file_path,search_start):
    def make_list(search_text, origin_text):
        sequence_list = []
        for search_text_elem in search_text:
            iter = re.finditer(search_text_elem, origin_text)
            for text in iter:
                start = text.start()
                end = text.end()
                # print(start,end)
                sequence_list.append([start, end])
        sequence_list.sort(key=lambda x: x[0])
        # print("리스트는:", sequence_list)
        new_list = sequence_list.copy()
        # print("카피된리스트는:", new_list)
        # print("요소갯수는:", len(sequence_list))
        for index, elem in enumerate(sequence_list):
            if index != len(sequence_list) - 1:
                if elem[1] >= sequence_list[index + 1][0]:
                    print("제거")
                    sequence_list[index] = [elem[0], max(elem[1], sequence_list[index + 1][1])]
                    sequence_list[index + 1] = [elem[0], max(elem[1], sequence_list[index + 1][1])]
        result = []
        for value in sequence_list:
            if value not in result:
                result.append(value)

        sequence_list = result
        # print("중복제거된리스트:", sequence_list)
        if sequence_list[0][0] == 0:
            odd_flag = False
        else:
            odd_flag = True

        # print("행열갯수는:", len(sequence_list))
        addon = []
        for i in range(0, len(sequence_list) + 1):
            addon.append(2 * i)
        # print(addon)
        for elem in addon:
            sequence_list.insert(elem, [0, 0])
        # print(sequence_list)
        for index, elem in enumerate(sequence_list):
            # print(elem,index)
            if elem[0] == 0 and elem[1] == 0:
                if index == 0:
                    elem[0] = 0
                    elem[1] = sequence_list[index + 1][0]
                elif index == len(sequence_list) - 1:
                    elem[0] = sequence_list[index - 1][1]
                    elem[1] = len(origin_text)
                else:
                    elem[0] = sequence_list[index - 1][1]
                    elem[1] = sequence_list[index + 1][0]
        # print(sequence_list)

        split_text = []
        for result in sequence_list:
            output = origin_text[result[0]:result[1]]
            # print(output)
            split_text.append(output)

        split_texts = list(filter(None, split_text))
        # print(split_texts)

        return split_texts, odd_flag

    # ------------------------- 텍스트들 가져오는 기능
    f = open(file_path, 'r', encoding='utf-8')
    rdr = csv.reader(f)
    data_csv = []
    for index, line in enumerate(rdr):
        if index == 0:
            continue
        # print(line)
        data_csv.append(line)
    f.close()

    # data_csv가 읽어온 데이타 전부
    # print("data_csv:",data_csv)
    no_row = len(data_csv)
    # print("행의수:", no_row)
    text_list = []
    for i in range(0, no_row):
        is_result = len(data_csv[i])
        # print("is_result:",is_result)
        if is_result >= 8:
            print("발췌문있다.{}번째".format(i))
            for j in range(7, is_result):
                print("j:",j,data_csv[i][j])

                text_list.append(
                    [i + 2, j + 1, data_csv[i][j], data_csv[i][0], data_csv[i][1], data_csv[i][2], data_csv[i][3],
                     data_csv[i][6]])

    # print("텍스트리스트:", text_list)



    # -------------------키워드들 가져오는 기능

    wb = openpyxl.load_workbook('keyword.xlsx')
    ws = wb.active
    no_row = ws.max_row
    no_col = ws.max_column
    # print("마지막행은:",no_row)
    keyword_list = []
    keyword_first=[]
    for j in range(2, no_col + 1):
        for i in range(2, no_row + 1):

            keyword = ws.cell(row=i, column=j).value
            if keyword == "" or keyword == None:
                continue
            keyword_list.append(keyword)
            if j==2:
                keyword_first.append(ws.cell(row=i, column=j).value)


    name = keyword_first[0]
    title = keyword_first[1]
    if title.find(",")>=0:
        # print("쉼표있음")
        title_list=title.split(",")
        # print("title_list:",title_list)
    else:
        # print("쉼표없음")
        title_list=[title]
    new_keyword_first = []
    for title_elem in title_list:
        tails = ['은', '는', '도',' 또한','이','가',' 역시']
        for tail in tails:
            result1 = name + tail
            result2 = name + " " + title_elem + tail
            result3 = name[0] + " " + title_elem + tail
            new_keyword_first.append(result1)
            new_keyword_first.append(result2)
            new_keyword_first.append(result3)
    keyword_first = new_keyword_first

    # additional_word=['그러면서','그는','또','이어','나아가','아울러','계속해서','또한']
    # keyword_first.extend(additional_word)
    keyword_first.append("“")
    keyword_first.append("”")
    keyword_first.append('"')
    keyword_first.append('"')
    keyword_first_only_for_search=keyword_first[:-4]
    # print("2.keyword_first_only_for_search:", keyword_first_only_for_search)
    # print("2.keyword_first:", keyword_first)


    doc = Document()
    for text in text_list:
        para = doc.add_paragraph("")
        search_text = keyword_list
        #---------------문단 색칠 쪼개기
        try:
            split_texts, odd_flag = make_list(search_text, text[2])
        except:
            # print("스프릿안됨")
            split_texts,odd_flag=make_list(search_start,text[2])

        if odd_flag == True:
            determinant = 1
        else:
            determinant = 0

        # -------------문장 단위 쪼개기
        sentence_each_list = []  # 문장 단위로 쪼개는 기능

        sentence_elem_split = text[2].split(".")
        sentence_each_list.extend(sentence_elem_split)
        # print('sentence_each_list:', sentence_each_list)

        new_one = []
        for sentence_each_elem in sentence_each_list:  # 문장단위로 쪼갠것에서 빈 요소 지우기
            print("test:", sentence_each_elem, len(sentence_each_elem))
            if len(sentence_each_elem) >= 1:
                new_one.append(sentence_each_elem)
        sentence_each_list = new_one
        # print('sentence_each_list:',sentence_each_list) #문장단위 출력하기

        sentence_search_list = []

        for sentence_each_elem in sentence_each_list:
            sentence_each_elem_origin=copy.copy(sentence_each_elem)
            if sentence_each_elem.find("“")>=0:
                position_fr = sentence_each_elem.find("“")
                position_rr = sentence_each_elem.find("”")
                if position_fr < 0 or position_rr < 0:
                    # print("따옴표 한쪽만 있어서 에러로 간주")
                    continue

                # print('keyword_first_first_only_for_serach:', keyword_first_only_for_search)
                # print("sentence_each_elem:", sentence_each_elem)

                for keyword_first_elem in keyword_first_only_for_search:
                    position_keyword = sentence_each_elem.find(keyword_first_elem)
                    # print(position_keyword,position_fr,position_rr,len(sentence_each_elem))
                    if 0 <= position_keyword < position_fr or position_rr < position_keyword < len(sentence_each_elem):
                        sentence_search_list.append(sentence_each_elem_origin)
            elif sentence_each_elem.find('"')>=0:
                regex = re.compile('"')
                cut_index=[]
                for elem in regex.finditer(sentence_each_elem):
                    cut_index.append(elem.start())
                # print(cut_index)
                result = cut_list(cut_index, 2)
                # print("cut_list", result)
                for elem in result:
                    if len(elem)==1:
                        result=result[:-1]
                        print("수정된cut_list", result)
                if len(result)==0:
                    continue
                text_list = []

                for elem in result:
                    text_cut = sentence_each_elem[elem[0]:elem[1] + 1]
                    text_list.append(text_cut)
                # print("따옴표안쪽텍스트:",text_list)
                inner_find=True
                for text_elem in text_list:
                    for keyword_first_elem in keyword_first_only_for_search:
                        if text_elem.find(keyword_first_elem)>=0:
                            inner_find=False
                # print("inner_find:",inner_find)

                for text_elem in text_list:
                    result = sentence_each_elem.replace(text_elem, "")
                    sentence_each_elem = result
                # print(result)

                outer_find = False
                for keyword_first_elem in keyword_first_only_for_search:
                    if result.find(keyword_first_elem)>=0:
                        outer_find=True
                # print('outer_find:',outer_find)
                if inner_find==True and outer_find==True:
                    sentence_search_list.append(sentence_each_elem_origin)



        # print('sentence_search_list:', sentence_search_list)

        #-------------------문장 색칠 쪼개기
        split_texts2=[]
        for sentence_search_elem in sentence_search_list:
            try:
                split_texts2, odd_flag2 = make_list(keyword_first, sentence_search_elem)
            except:
                print("스프릿안됨")
                split_texts2,odd_flag2=make_list(keyword_first,sentence_search_elem)

            if odd_flag2 == True:
                determinant2 = 1
            else:
                determinant2 = 0

        para.add_run("1.행번호 : ")
        para.add_run(str(text[0]))
        para.add_run("\n")
        para.add_run("2.열번호 : ")
        para.add_run(str(text[1]))
        para.add_run("\n")
        para.add_run("3.URL : ")
        para.add_run(str(text[3]))
        para.add_run("\n")
        para.add_run("4.신문사 : ")
        para.add_run(str(text[4]))
        para.add_run("\n")
        para.add_run("5.제목 : ")
        para.add_run(str(text[5]))
        para.add_run("\n")
        para.add_run("6.날짜 : ")
        para.add_run(str(text[6]))
        para.add_run("\n")
        para.add_run("7.기자 : ")
        para.add_run(str(text[7]))
        para.add_run("\n")
        para.add_run("8.발췌문단 : ")
        para.add_run("\n")

        for index, split_text in enumerate(split_texts):

            if index % 2 == determinant:
                para.add_run(split_text)
            else:
                para.add_run(split_text)

        para.add_run("\n")
        para.add_run("9.발췌문장 : ")
        para.add_run("\n")
        for index, split_text2 in enumerate(split_texts2):
            if index % 2 == determinant2:
                para.add_run(split_text2)
            else:
                para.add_run(split_text2)
        # for sentence_search_elem in sentence_search_list:
        #     para.add_run(str(sentence_search_elem))
        #     para.add_run("\n")

        doc.add_page_break()
        para2 = doc.add_paragraph("")


    doc.save(file_path.replace('.csv', '.docx'))
    # ----------------------------------------
def get_list(keyword,start_date,end_date,page_no):
    cookies = {
        'Bigkinds': 'BCA046FB36C525CF281497A3841C3BCB',
        '_ga': 'GA1.3.1113371620.1679298308',
        '_gid': 'GA1.3.1569985393.1679298308',
    }

    headers = {
        'Accept': 'application/json, text/javascript, */*; q=0.01',
        'Accept-Language': 'ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7',
        'Connection': 'keep-alive',
        'Content-Type': 'application/json;charset=UTF-8',
        # 'Cookie': 'Bigkinds=BCA046FB36C525CF281497A3841C3BCB; _ga=GA1.3.1113371620.1679298308; _gid=GA1.3.1569985393.1679298308',
        'Origin': 'https://www.bigkinds.or.kr',
        'Referer': 'https://www.bigkinds.or.kr/v2/news/index.do',
        'Sec-Fetch-Dest': 'empty',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Site': 'same-origin',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/111.0.0.0 Safari/537.36',
        'X-Requested-With': 'XMLHttpRequest',
        'sec-ch-ua': '"Google Chrome";v="111", "Not(A:Brand";v="8", "Chromium";v="111"',
        'sec-ch-ua-mobile': '?0',
        'sec-ch-ua-platform': '"Windows"',
    }

    json_data = {
        'indexName': 'news',
        'searchKey': keyword,
        'searchKeys': [
            {},
        ],
        'byLine': '',
        'searchFilterType': '3',
        'searchScopeType': '1',
        'searchSortType': 'date',
        'sortMethod': 'date',
        'mainTodayPersonYn': '',
        'startDate': start_date,
        'endDate': end_date,
        'newsIds': None,
        'categoryCodes': [],
        'providerCodes': [
            '01100101',
            '01100201',
            '01100301',
            '01100401',
            '01100501',
            '01100611',
            '01100701',
            '01100801',
            '01100901',
            '01101001',
            '01101101',
        ],
        'incidentCodes': [],
        'networkNodeType': '',
        'topicOrigin': '',
        'dateCodes': [],
        'editorialIs': False,
        'startNo': page_no,
        'resultNumber': '100',
        'isTmUsable': False,
        'isNotTmUsable': False,
    }

    response = requests.post('https://www.bigkinds.or.kr/api/news/search.do', cookies=cookies, headers=headers,
                             json=json_data)
    # print(response.text)
    return response
    # Note: json_data will not be serialized by requests
    # exactly as it was in the original request.
    #data = '{"indexName":"news","searchKey":"","searchKeys":[{}],"byLine":"","searchFilterType":"1","searchScopeType":"1","searchSortType":"date","sortMethod":"date","mainTodayPersonYn":"","startDate":"2022-12-20","endDate":"2023-03-20","newsIds":[],"categoryCodes":[],"providerCodes":[],"incidentCodes":[],"networkNodeType":"","topicOrigin":"","dateCodes":[],"editorialIs":false,"startNo":1,"resultNumber":10,"isTmUsable":false,"isNotTmUsable":false}'
    #response = requests.post('https://www.bigkinds.or.kr/api/news/search.do', cookies=cookies, headers=headers, data=data)
def get_detail(article_no):
    import requests

    cookies = {
        'Bigkinds': 'BCA046FB36C525CF281497A3841C3BCB',
        '_ga': 'GA1.3.1113371620.1679298308',
        '_gid': 'GA1.3.1569985393.1679298308',
    }

    headers = {
        'Accept': 'application/json, text/javascript, */*; q=0.01',
        'Accept-Language': 'ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7',
        'Connection': 'keep-alive',
        # 'Cookie': 'Bigkinds=BCA046FB36C525CF281497A3841C3BCB; _ga=GA1.3.1113371620.1679298308; _gid=GA1.3.1569985393.1679298308',
        'Referer': 'https://www.bigkinds.or.kr/v2/news/index.do',
        'Sec-Fetch-Dest': 'empty',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Site': 'same-origin',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/111.0.0.0 Safari/537.36',
        'X-Requested-With': 'XMLHttpRequest',
        'sec-ch-ua': '"Google Chrome";v="111", "Not(A:Brand";v="8", "Chromium";v="111"',
        'sec-ch-ua-mobile': '?0',
        'sec-ch-ua-platform': '"Windows"',
    }

    params = {
        'docId': article_no,
        'returnCnt': '1',
        'sectionDiv': '1000',
    }

    response = requests.get('https://www.bigkinds.or.kr/news/detailView.do', params=params, cookies=cookies,
                            headers=headers)
    return response
def cut_list(my_list,n):
    result = [my_list[i * n:(i + 1) * n] for i in range((len(my_list) + n - 1) // n )]
    return result
def get_keyword():
    wb=openpyxl.load_workbook('keyword.xlsx')
    ws=wb.active
    name=ws.cell(row=2,column=2).value
    position=ws.cell(row=3,column=2).value
    position_list=position.split(",")
    add_on_list=['은','는','도','이','가',' 역시']
    keyword_list=[]
    for i in range(0,3): # 성 분화
        if i==1:
            name = name + "(사진)"
        if i==2:
            name=name[0]
        for position_elem in position_list: #직책 분화
            position_elem_garo=name+" "+position_elem+"("
            keyword_list.append(position_elem_garo)
            for add_on_elem in add_on_list: #토씨분화
                name_add_on=name+" "+position_elem+add_on_elem
                keyword_list.append(name_add_on)

    print('keyword_list:',keyword_list)
    return keyword_list

def get_passage_list(fname):
    # CSV파일에서 글 각각 가져오기
    origin_passage_list=[]
    j = open(fname, 'r', encoding='utf-8-sig', newline="")
    rdr = csv.reader(j)
    for index, line in enumerate(rdr):
        if index == 0:
            continue
        data=[index+1,line[0],line[1],line[2],line[3],line[4].replace('“','"').replace('”','"').replace('""','"').replace("\t",""),"",line[5]]
        # print(data)
        origin_passage_list.append(data)
    j.close()
    return origin_passage_list

def find_sentence_passage(origin_passage_list,keyword_list,start_word_list,explanation_list,additional_word_list):
    #한개의 글에 대해서 문장으로 나누기
    search_result_total=[] # 전체 검색 결과
    count = 0
    for origin_passage_elem in origin_passage_list:
        # print('origin_passage_elem[5]:',origin_passage_elem[5])

        ddaomb_group=[]
        ddaomb_set = []

        #숫자 안에 있는 점은 다른걸로 바꿈
        regex = re.compile("\d+.\d+")
        for regex_elem in regex.finditer(origin_passage_elem[5]):
            changed_regex = regex_elem.group().replace(".", '#')
            # print(changed_regex)
            origin_passage_elem[5] = origin_passage_elem[5].replace(regex_elem.group(), changed_regex)


        for index,i in enumerate(origin_passage_elem[5]): #따옴표 셋트 위치 찾기
            if i=='"':
                ddaomb_set.append(index)
            if len(ddaomb_set)==2:
                ddaomb_group.append(ddaomb_set)
                ddaomb_set=[]
        # print('ddaomb_group:',ddaomb_group)

        position_point=[]
        for index,i in enumerate(origin_passage_elem[5]):
            if i==".":
                # print('index:',index)
                position_point.append(index)

        for position_point_elem in position_point:
            for ddaomb_group_elem in ddaomb_group:
                if ddaomb_group_elem[0]<=position_point_elem<=ddaomb_group_elem[1]:

                    # print("따옴표 사이 점 슬래시로 치환",position_point_elem)
                    # print(origin_passage_elem[5][position_point_elem])
                    new_origin_passage_elem=list(origin_passage_elem[5])
                    new_origin_passage_elem[position_point_elem]="/"
                    origin_passage_elem[5]="".join(new_origin_passage_elem)


        sentence_list=origin_passage_elem[5].split(".")

        #각 문장에서 빈거 없애고, 양쪽에 공백 없애서 다시 넣기
        new_sentence_list=[]
        for sentence_elem in sentence_list:
            if len(sentence_elem)>=2:
                sentence_elem=sentence_elem.strip()
                new_sentence_list.append(sentence_elem)
        sentence_list=new_sentence_list
        # print('sentence_list:',sentence_list)

        #문단만들기
        sentence_group=origin_passage_elem[5].split("\n")
        new_sentence_group=[]
        for sentence_group_elem in sentence_group:
            if len(sentence_group)>=0:
                new_sentence_group.append(sentence_group_elem)
        sentence_group=new_sentence_group
        # print('sentence_group:',sentence_group)


        #기존에 충족했는지 여부를 저장
        check_deque=deque([0,0,0,0,0],maxlen=5)
        test=[]

        for index_sentence_elem,sentence_elem in enumerate(sentence_list):
            # print('sentence_elem:',sentence_elem)
            check_index=0 #검출됐는지를 체크하는 인덱스
            search_result_sentence=""
            regex=re.compile('"')
            regex.finditer(sentence_elem)


            spliter_list=[]
            for regex_elem in regex.finditer(sentence_elem):
                # print(regex_elem.start())
                spliter_list.append(regex_elem.start())
            # print("spliter_length:",len(spliter_list))

            spliter_list = []
            spliter_set = []

            for index_spliter,spliter in enumerate(regex.finditer(sentence_elem)):
                # print(spliter.start())
                if index_spliter%2==0:
                    spliter_set.append(spliter.start())
                else:
                    spliter_set.append(spliter.start()+1)

                if len(spliter_set)==0:
                    # print("빈행렬")
                    pass
                elif len(spliter_set)==2:
                    # print("찬행렬")
                    pass
                    spliter_list.append(spliter_set)
                    spliter_set = []
            # print('spliter_list:',spliter_list)
            sentence_inner_all=""
            sentence_inner_list=[]
            sentence_outer_all=""
            for i in range(0,len(spliter_list)):
                for spliter_elem in spliter_list:
                    if i==0:
                        sentence_inner=sentence_elem[spliter_elem[0]:spliter_elem[1]]
                        sentence_inner_list.append(sentence_inner)
                        sentence_inner_all=sentence_inner_all+sentence_inner

                    sentence_outer_all=sentence_elem
                    for sentence_inner_elem in sentence_inner_list:
                        sentence_outer_all=sentence_outer_all.replace(sentence_inner_elem,"%%")


            #★★★★★★★★★★★★★★★★★★★★기준에 충족하는지 여부 확인 하는 부분★★★★★★★★★★★★★★★★★★★★★★★★★★★★★★★★
            check_index=0
            print("sentence_elem:",sentence_elem)
            # print('keyword_list:',keyword_list)
            status=""

            for index,keyword_elem in enumerate(keyword_list): # 안쪽과 바깥쪽에 단어 여부 확인
                # if sentence_inner_all.find(keyword_elem)<0 and sentence_outer_all.find(keyword_elem)>=0:
                if sentence_inner_all.find(keyword_elem)<0 and sentence_outer_all.find(keyword_elem)>=0 :
                    text='케이스1의 문장'
                    print(text)
                    status=status+text+"\n"
                    # print('sentence_inner_all:',sentence_inner_all)
                    # print('new_sentence_outer_all:',sentence_outer_all)
                    search_result_sentence=sentence_elem
                    check_index=1



            for keyword_elem in keyword_list:  # 안쪽과 바깥쪽에 단어 여부 확인
                # print(keyword_elem)
                for explanation_elem in explanation_list:
                    if sentence_inner_all.find(keyword_elem)<0 and (sentence_outer_all.find(keyword_elem)>=0 and sentence_outer_all.find(keyword_elem)<sentence_outer_all.find("%%")) and sentence_outer_all.find(explanation_elem)>=0 and sentence_outer_all.find(explanation_elem)>sentence_outer_all.find("%%"):
                        text='케이스2의 문장'
                        print(text)
                        status = status + text+"\n"
                        # print('sentence_inner_all:',sentence_inner_all)
                        # print('new_sentence_outer_all:',sentence_outer_all)
                        search_result_sentence=sentence_elem
                        check_index = 2




            # 핵심키워드말고 추가 서술에도 고려하여 검색
            search_all_list = start_word_list+additional_word_list
            if check_deque[-1]==1 or check_deque[-1]==2:
                for search_all_elem in search_all_list:
                    if sentence_outer_all.find(search_all_elem)>=0  and sentence_outer_all.find(search_all_elem)<sentence_outer_all.find("%%"):
                        text='케이스3-1의 문장 {}'.format(check_deque)
                        print(text)
                        status = status + text+"\n"
                        # print('sentence_inner_all:',sentence_inner_all)
                        # print('new_sentence_outer_all:',sentence_outer_all)
                        search_result_sentence=sentence_elem
                        check_index = 3


            search_all_list = start_word_list + additional_word_list
            if check_deque[-1]==1 or check_deque[-1]==2:
                for search_all_elem in search_all_list:
                    for explanation_elem in explanation_list:
                        if sentence_outer_all.find(search_all_elem)>=0 and sentence_outer_all.find(search_all_elem)<sentence_outer_all.find("%%")  and sentence_outer_all.find(explanation_elem)>=0 and sentence_outer_all.find(explanation_elem)>sentence_outer_all.find("%%"):
                            text='케이스3-2의 문장 {}'.format(check_deque)
                            print(text)
                            status = status + text+"\n"
                            # print('sentence_inner_all:',sentence_inner_all)
                            # print('new_sentence_outer_all:',sentence_outer_all)
                            search_result_sentence=sentence_elem
                            check_index = 3

            if check_deque[-1]==1 or check_deque[-1]==2:
                if sentence_outer_all.startswith('"') or sentence_outer_all.startswith('그러나 "') :
                    text='케이스3-3의 문장 {}'.format(check_deque)
                    print(text)
                    status = status + text+"\n"
                    # print('sentence_inner_all:',sentence_inner_all)
                    # print('new_sentence_outer_all:',sentence_outer_all)
                    search_result_sentence=sentence_elem
                    check_index = 3


            search_all_list = start_word_list + additional_word_list
            for keyword_elem in keyword_list:
                if sentence_list[index_sentence_elem-1].find(keyword_elem)>=0:
                    for search_all_elem in search_all_list:
                            if sentence_outer_all.find(search_all_elem) >= 0 and sentence_outer_all.find(search_all_elem)<sentence_outer_all.find("%%"):
                                text='케이스4-1의 문장 {}'.format(check_deque)
                                print(text)
                                status = status + text+"\n"
                                # print('sentence_inner_all:',sentence_inner_all)
                                # print('new_sentence_outer_all:',sentence_outer_all)
                                search_result_sentence = sentence_elem
                                check_index = 4

            search_all_list = start_word_list + additional_word_list
            for keyword_elem in keyword_list:
                if sentence_list[index_sentence_elem-1].find(keyword_elem)>=0:
                    for search_all_elem in search_all_list:
                        for explanation_elem in explanation_list:
                            if sentence_outer_all.find(search_all_elem) >= 0 and sentence_outer_all.find(search_all_elem)<sentence_outer_all.find("%%") and sentence_outer_all.find(explanation_elem) >= 0 and sentence_outer_all.find(explanation_elem)>sentence_outer_all.find("%%"):
                                text='케이스4-2의 문장 {}'.format(check_deque)
                                print(text)
                                status = status + text+"\n"
                                # print('sentence_inner_all:',sentence_inner_all)
                                # print('new_sentence_outer_all:',sentence_outer_all)
                                search_result_sentence = sentence_elem
                                check_index = 4


            for keyword_elem in keyword_list:
                if sentence_list[index_sentence_elem-1].find(keyword_elem)>=0:
                        if sentence_elem.startswith('"'):
                            text='케이스4-3의 문장 {}'.format(check_deque)
                            print(text)
                            status = status + text+"\n"
                            # print('sentence_inner_all:',sentence_inner_all)
                            # print('new_sentence_outer_all:',sentence_outer_all)
                            search_result_sentence = sentence_elem
                            check_index = 4

            search_all_list = start_word_list+additional_word_list
            if check_deque[-2] == 1 or check_deque[-2] == 2:
                if check_deque[-1]==3:
                    for search_all_elem in search_all_list:
                        if sentence_outer_all.find(search_all_elem)>=0 and sentence_outer_all.find(search_all_elem)<sentence_outer_all.find("%%"):
                            text='케이스5-1의 문장 {}'.format(check_deque)
                            print(text)
                            status = status + text+"\n"
                            # print('sentence_inner_all:',sentence_inner_all)
                            # print('new_sentence_outer_all:',sentence_outer_all)
                            search_result_sentence=sentence_elem
                            check_index = 5

            search_all_list = start_word_list + additional_word_list+explanation_list
            if check_deque[-2] == 1 or check_deque[-2] == 2:
                if check_deque[-1] == 3:
                    for search_all_elem in search_all_list:
                        for explanation_elem in explanation_list:
                            if sentence_outer_all.find(search_all_elem) >= 0 and sentence_outer_all.find(search_all_elem)<sentence_outer_all.find("%%") and sentence_outer_all.find(explanation_elem)>=0 and sentence_outer_all.find(explanation_elem)>sentence_outer_all.find("%%"):
                                text='케이스5-2의 문장 {}'.format(check_deque)
                                print('케이스5-2의 문장 {}'.format(check_deque))
                                status = status + text+"\n"
                                # print('sentence_inner_all:',sentence_inner_all)
                                # print('new_sentence_outer_all:',sentence_outer_all)
                                search_result_sentence = sentence_elem
                                check_index = 5

            search_all_list = start_word_list+additional_word_list
            for keyword_elem in keyword_list:
                if sentence_list[index_sentence_elem - 2].find(keyword_elem) >= 0:
                    if check_deque[-1]==4:
                        for search_all_elem in search_all_list:
                            if sentence_outer_all.find(search_all_elem)>=0 and sentence_outer_all.find(search_all_elem)<sentence_outer_all.find("%%") :
                                text='케이스5-3의 문장 {}'.format(check_deque)
                                print(text)
                                status = status + text+"\n"
                                # print('sentence_inner_all:',sentence_inner_all)
                                # print('new_sentence_outer_all:',sentence_outer_all)
                                search_result_sentence=sentence_elem
                                check_index = 5

            search_all_list = start_word_list + additional_word_list+explanation_list
            for keyword_elem in keyword_list:
                if sentence_list[index_sentence_elem - 2].find(keyword_elem) >= 0:
                    if check_deque[-1] == 4:
                        for search_all_elem in search_all_list:
                            for explanation_elem in explanation_list:
                                if sentence_outer_all.find(search_all_elem) >= 0 and sentence_outer_all.find(search_all_elem)<sentence_outer_all.find("%%") and sentence_outer_all.find(explanation_elem)>=0 and sentence_outer_all.find(explanation_elem)>sentence_outer_all.find("%%"):
                                    text='케이스5-4의 문장 {}'.format(check_deque)
                                    print(text)
                                    status = status + text+"\n"
                                    # print('sentence_inner_all:',sentence_inner_all)
                                    # print('new_sentence_outer_all:',sentence_outer_all)
                                    search_result_sentence = sentence_elem
                                    check_index = 5


            search_result_passage=""
            for sentence_group_elem in sentence_group:
                if sentence_group_elem.find(search_result_sentence)>=0:
                    search_result_passage=sentence_group_elem

            check_deque.append(check_index) #현재 선택됐는지 여부를 데크에 저장
            if len(search_result_sentence)>=1:
                # print("★★★★★★★★★★★★★★★★★★★★★★★★★★★")
                # print('search_result:',search_result_sentence)
                # print('search_passage:',search_result_passage)
                # print("check_deque:",check_deque)
                # data = [index, line[0], line[1], line[2], line[3], line[4].replace('“', '"').replace('”', '"').replace('""', '"').replace("\n", "").replace("\t",""),line[5], line[6]]
                # test.append(search_result_sentence)

                data=[origin_passage_elem[0],origin_passage_elem[1],origin_passage_elem[2],origin_passage_elem[3],origin_passage_elem[4],origin_passage_elem[7],search_result_passage.replace("#","."),search_result_sentence.replace("#","."),status]
                print('count:',count)
                count=count+1
                search_result_total.append(data)
                print("====================================")

    return search_result_total

def split_text(sentence,position_keyword_color):
    split_index=[]
    color_index=0
    if position_keyword_color[0]==0:
        data=[position_keyword_color[0],position_keyword_color[1]]
        split_index.append(data)
        data=[position_keyword_color[1],len(sentence)]
        split_index.append(data)
        color_index=0
    elif position_keyword_color[1]==len(position_keyword_color):
        data=[0,position_keyword_color[0]]
        split_index.append(data)
        data=[position_keyword_color[0],position_keyword_color[1]]
        split_index.append(data)
        color_index=1
    else:
        data=[0,position_keyword_color[0]]
        split_index.append(data)
        data=[position_keyword_color[0],position_keyword_color[1]]
        split_index.append(data)
        data = [position_keyword_color[1], len(sentence)]
        split_index.append(data)
        color_index = 1
    print('split_index:',split_index)

    split_result_list=[]
    for split_index_elem in split_index:
        split_result_list.append(sentence[split_index_elem[0]:split_index_elem[1]])
    print('split_result_list:',split_result_list)
    return split_result_list,color_index
def split_text_advanced(text,keyword_list):
    nums=[]
    all_nums=[]
    for i in range(0,len(text)+1):
        all_nums.append(i)
    # print('all_nums:',all_nums)
    for keyword_elem in keyword_list:

        result=text.find(keyword_elem)
        if result>=0:
            keyword_length=len(keyword_elem)
            for i in range(0,keyword_length+1):
                # print(result+keyword_length)
                nums.append(result+i)

    # print("nums:",nums)
    rest_nums=[]
    for all_num in all_nums:
        if all_num not in nums:
            rest_nums.append(all_num)
    # print('rest_num:',rest_nums)

    def grouping(nums):
        sequences = []
        # 리스트의 첫 번째 요소를 시작으로 반복문을 실행합니다.
        start = nums[0]
        for i in range(1, len(nums)):
            # 이전 숫자와 현재 숫자가 연속된 숫자인 경우
            if nums[i] == nums[i-1]+1:
                # 현재 숫자를 계속해서 연속된 숫자로 묶어줍니다.
                continue
            else:
                # 연속된 숫자들을 리스트에 저장합니다.
                sequences.append([start, nums[i-1]])
                # 다음 연속된 숫자들을 찾기 위해 start 변수를 현재 숫자로 갱신합니다.
                start = nums[i]

        # 마지막으로 연속된 숫자들을 리스트에 저장합니다.
        sequences.append([start, nums[-1]])
        sequences.sort(key=lambda x: x[0])
        # print(sequences)
        return sequences
    new_nums=grouping(nums)
    new_rest_nums=grouping(rest_nums)
    # for new_num in new_nums:
    #     new_num[1]=new_num[1]+1

    # print('new_nums:',new_nums)
    # print('new_rest_nums:',new_rest_nums)
    total_nums=new_nums+new_rest_nums
    total_nums.sort(key=lambda x: x[0])
    for total_num in total_nums:
        total_num[1]=total_num[1]+1
    # print(total_nums)

    split_text_list=[]
    for total_num in total_nums:
        split_text=text[total_num[0]:total_num[1]]
        # print(split_text)
        split_text_list.append(split_text)


    if new_nums[0][0]==0:
        odd_even_flag="even"
    else:
        odd_even_flag = "odd"
    # print('split_text_list:',split_text_list)
    # print('odd_even_flag:',odd_even_flag)
    return split_text_list,odd_even_flag


class Thread(QThread):
    cnt = 0
    user_signal = pyqtSignal(str)  # 사용자 정의 시그널 2 생성
    user_signal2 = pyqtSignal()  # 사용자 정의 시그널 2 생성

    def __init__(self, parent,keyword,total_start_date,total_end_date):  # parent는 WndowClass에서 전달하는 self이다.(WidnowClass의 인스턴스)
        super().__init__(parent)
        self.parent = parent  # self.parent를 사용하여 WindowClass 위젯을 제어할 수 있다.
        self.keyword=keyword
        self.total_start_date=total_start_date
        self.total_end_date=total_end_date

    def run(self):
        datetime_now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        f = open('RESULT_{}_{}.csv'.format(self.keyword, datetime_now), 'w', encoding='utf-8-sig', newline="")
        wr = csv.writer(f)
        wr.writerow(['URL', "신문사", "제목", "날짜", "전문", "기자"])

        for num_page in range(1,9999):
            text="{}번째 페이지 크롤링 중...".format(num_page)
            self.user_signal.emit(text)
            print(num_page,"번째 페이지 이동")
            response = get_list(self.keyword,self.total_start_date,self.total_end_date,num_page)
            result = json.loads(response.text)
            result_details = result['resultList']
            if len(result_details)==0:
                print("페이지 더 없음")
                break
            for index, result_detail in enumerate(result_details):
                title = result_detail['TITLE']
                news_id = result_detail['NEWS_ID']


                response = get_detail(news_id)
                json_detail = json.loads(response.text)
                # pprint.pprint(json_detail)
                contents = json_detail['detail']['CONTENT'].replace("<br/>","\n")
                provider = json_detail['detail']['PROVIDER']
                title = json_detail['detail']['TITLE']
                try:
                    image_url = json_detail['detail']['IMAGES']
                    if image_url.find(",") >= 0:
                        image_url_list = image_url.split(",")
                        image_url_list = str(image_url_list)
                    else:
                        image_url_list = image_url
                except:
                    image_url = ""
                news_date = json_detail['detail']['DATE']
                reporter = json_detail['detail']['BYLINE']
                provider_url=json_detail['detail']['PROVIDER_LINK_PAGE']
                print("{}번째 크롤링 결과".format(index))
                print("신문사:", provider)
                print("타이틀:", title)
                print("내용:", contents)
                print("리포터:", reporter)
                print("이미지URL:", image_url_list)
                print("----------------")
                time.sleep(0.5)
                data=[provider_url,provider,title,news_date,contents,reporter]
                wr.writerow(data)

        f.close()
        self.user_signal2.emit()
    def stop(self):
        pass


class Example(QMainWindow,Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.path="C:"
        self.index=None
        self.setupUi(self)
        self.setSlot()
        self.show()

        self.dateEdit_2.setDate(QDate.currentDate())
        time_1month_ago=datetime.datetime.now()-datetime.timedelta(days=30)
        time_1month_ago_year=int(time_1month_ago.strftime("%Y"))
        time_1month_ago_month = int(time_1month_ago.strftime("%m"))
        time_1month_ago_day = int(time_1month_ago.strftime("%d"))
        self.dateEdit.setDate(QDate(time_1month_ago_year,time_1month_ago_month,time_1month_ago_day))
        self.auth_flag=True # 테스트용 auth always True code
        # self.auth_flag = False
        self.first_flag=True
        self.lineEdit_5.setPlaceholderText("프로그램 번호를 입력하세요")

    def start(self):

        if self.auth_flag==True:
            print(" 로그인성공")
        elif self.auth_flag==False:
            QMessageBox.information(self, "실패창", "인증 후에 프로그램을 실행하여 주십시오.")
            return

        self.keyword = self.lineEdit_2.text()
        self.start_date_year=str(self.dateEdit.date().year())
        self.start_date_month = str(self.dateEdit.date().month())
        if len(self.start_date_month)==1:
            self.start_date_month="0"+self.start_date_month
        self.start_date_day = str(self.dateEdit.date().day())
        if len(self.start_date_day)==1:
            self.start_date_day="0"+self.start_date_day
        self.total_start_date=self.start_date_year+"-"+self.start_date_month+"-"+self.start_date_day
        print("시작일:",self.total_start_date)
        self.end_date_year = str(self.dateEdit_2.date().year())
        self.end_date_month = str(self.dateEdit_2.date().month())
        if len(self.end_date_month) == 1:
            self.end_date_month = "0" + self.end_date_month
        self.end_date_day = str(self.dateEdit_2.date().day())
        if len(self.end_date_day) == 1:
            self.end_date_day = "0" + self.end_date_day
        self.total_end_date = self.end_date_year+"-"+ self.end_date_month+"-"+self.end_date_day
        print("종료일:",self.total_end_date)
        print("키워드는:",self.keyword)



        self.x = Thread(self,self.keyword,self.total_start_date,self.total_end_date)
        self.x.user_signal.connect(self.slot1)  # 사용자 정의 시그널2 슬롯 Connect
        self.x.user_signal2.connect(self.slot2)  # 사용자 정의 시그널2 슬롯 Connect
        self.x.start()
    def slot1(self,data1):
        self.textEdit.append(str(data1))

    def slot2(self):
        QMessageBox.information(self, "완료창", "작업이 완료 되었습니다.")

    def auth(self):
        text=str(self.lineEdit_5.text())
        print(text)
        result_password,result_ip = get_key("pc{}".format(text),self.first_flag)
        print('나의_IP:',result_ip,'나의_PASSWORD:',result_password)
        self.first_flag=False
        # print(result)
        self.password=self.lineEdit_4.text()
        # print('지정_IP:',socket.gethostbyname(socket.gethostname()),'현재_Password:',self.password)
        print('지정된_IP:', socket.gethostbyname(socket.gethostname()))
        if self.password==result_password and socket.gethostbyname(socket.gethostname())==result_ip:
            print("로그인 완료")
            self.auth_flag=True
            QMessageBox.information(self, "로그인", "인증이 성공하였습니다.")
        else:
            print("로그인 실패")
            self.auth_flag=False
            QMessageBox.information(self, "로그인", "인증이 실패하였습니다.")



    def setSlot(self):
        pass

    def slot_fileopen(self):
        pass
    def setIndex(self, index):
        pass

    def quit(self):
        QCoreApplication.instance().quit()
    def find(self):
        if self.auth_flag==True:
            print(" 로그인성공")
        elif self.auth_flag==False:
            QMessageBox.information(self, "실패창", "인증 후에 프로그램을 실행하여 주십시오.")
            return
        print("find")
        self.fname=QFileDialog.getOpenFileName(self,"Open file",'./')[0]
        print(self.fname)
        self.lineEdit.setText(self.fname)

    def search(self):
        if self.auth_flag==True:
            print(" 로그인성공")
        elif self.auth_flag==False:
            QMessageBox.information(self, "실패창", "인증 후에 프로그램을 실행하여 주십시오.")
            return

        keyword_list = get_keyword()
        start_word_list = ['그러면서', '그는', '그들은', '이들은', '또 ', '이어', '나아가', '아울러', '계속해서', '또한','다만','즉','따라서','이에','그러면서도']
        explanation_list = ['고 내다봤', '고 논평', '고 답변', '고 답했', '고 덧붙', '고 반박', '고 밝혔', '고 비판', '고 설명', '고 썼', '고 압박',
                            '고 언급', '을 언급', '고 역설', '고 예상', '고 요구',
                            '고 우려', '고 일갈', '고 일축', '고 적었', '고 전했', '고 주장', '고 지적', '고 직격', ' 고 진화', '고 촉구', '고 해석',
                            '고 강하게 비판', '고 했', '고 말했', '고 말하기도', '고 강조', '고 거들었', '고 경고', '고 공격', '고 반문','고 꼬집었',
                            '며 꼬집었','고 당부','고 되물었','고 따졌','고 따져물었','며 말을 아꼈','면서도 말을 아꼇','고 맹공격','고 목소리를 높였',
                            '며 물러서지 않았','고 물었','고 반격','고 반문','반응을 보였','고 밝히','고 보탰','고 선을 그었','고 쏘아붙였','고 쓴소리',
                            '고 약속','고 역공','고 진단','고 질의','고 평가','고 표출','고 두둔','고 직격탄을 날렸','고 질타','며 즉답을 피했',
                            '고 추궁','고 추측','고 해명','고 호소','고도 했']
        additional_word_list = [' 대해선', ' 대해서는', ' 대해서도', ' 관해', ' 관해선', ' 관해서는', ' 관해서도', ' 관련해서는', ' 관련해선',
                                ' 묻자', ' 향해서는', ' 지적엔', ' 지적에는', ' 지적에도', ' 질문에는', ' 질문에', ' 질문엔', ' 질문에도',' 물음에는',
                                ' 물음엔',' 지적하자',' 지적에는',' 질의에는',' 질의엔',' 조언으로는',' 두고는',' 두곤',' 방안으로는',' 시각에는',' 년에는']
        origin_passage_list = get_passage_list(self.fname)
        # print('origin_passage_list:',origin_passage_list)
        search_result_total = find_sentence_passage(origin_passage_list,keyword_list,start_word_list,explanation_list,additional_word_list)



        doc = Document()
        for search_result_elem in search_result_total:
            # print('search_result_elem:',search_result_elem)


            keyword_list_start_word_list = keyword_list + start_word_list+ explanation_list+additional_word_list
            # print('keyword_list_start_word_list:', keyword_list_start_word_list)

            position_keyword_color = []
            for keyword_list_start_word_elem in keyword_list_start_word_list:
                position = search_result_elem[-1].find(keyword_list_start_word_elem)
                if position >= 0:
                    data = [position, position + len(keyword_list_start_word_elem)]
                    position_keyword_color = data
                    break

            # print('position_keyword_color:', position_keyword_color)


            para = doc.add_paragraph("")
            para.add_run("1.행번호 : ")
            para.add_run(str(search_result_elem[0]))
            para.add_run("\n")
            para.add_run("2.URL : ")
            para.add_run(str(search_result_elem[1]))
            para.add_run("\n")
            para.add_run("3.신문사 : ")
            para.add_run(str(search_result_elem[2]))
            para.add_run("\n")
            para.add_run("4.제목 : ")
            para.add_run(str(search_result_elem[3]))
            para.add_run("\n")
            para.add_run("5.날짜 : ")
            para.add_run(str(search_result_elem[4]))
            para.add_run("\n")
            para.add_run("6.기자 : ")
            para.add_run(str(search_result_elem[5]))
            para.add_run("\n")
            para.add_run("7.발췌문단 : ")
            para.add_run("\n")
            para.add_run(str(search_result_elem[6]))
            para.add_run("\n")
            para.add_run("8.발췌문장 : ")
            para.add_run("\n")
            # para.add_run(str(search_result_elem[-1]))
            if len(position_keyword_color)==0:
                para.add_run(search_result_elem[7])
            else:
                # split_result_list, color_index = split_text(search_result_elem[-1], position_keyword_color)
                # for index, split_result_elem in enumerate(split_result_list):
                #     if split_result_elem.find("/")>=0:
                #         split_result_elem.replace("/",".")
                #     if index == color_index:
                #         para.add_run(split_result_elem).font.color.rgb = RGBColor(0xFF, 0x24, 0xE9)
                #     else:
                #         para.add_run(split_result_elem)
                split_result_list, even_odd_flag = split_text_advanced(search_result_elem[7].replace("/","."), keyword_list_start_word_list)
                for index, split_result_elem in enumerate(split_result_list):
                    if even_odd_flag=='even':
                        if index%2==0:
                            para.add_run(split_result_elem)
                        else:
                            para.add_run(split_result_elem)
                    else:
                        if index%2==1:
                            para.add_run(split_result_elem)
                        else:
                            para.add_run(split_result_elem)
            para.add_run("\n")
            para.add_run("9.발췌로직 : ")
            para.add_run("\n")
            para.add_run(str(search_result_elem[8]))

            doc.add_page_break()
        print("저장하기")
        # doc.save('result.docx')
        time_now=datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        doc.save(self.fname.replace(".csv",'_SEARCH_RESULT_{}.docx'.format(time_now)))


        QMessageBox.information(self, "완료창", "작업이 완료 되었습니다.")


app=QApplication([])
ex=Example()
sys.exit(app.exec_())



if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = MyApp()
    sys.exit(app.exec_())